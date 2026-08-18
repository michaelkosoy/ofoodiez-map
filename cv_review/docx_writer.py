"""Generate the optimized CV as a brand-new .docx from the canonical
structured model (python-docx). NOTHING from the uploaded file is copied —
no XML parts, styles, relationships, macros, embedded objects or images can
survive into this document because it never touches the original package.
Links appear as plain text (no external relationships at all), which keeps
the final package inspection (docx_inspect.py) trivially clean.

Template mirrors the PDF: centered name + contact line, ruled "Experience:"
headings, a narrow left date column per position, tight bullets with bolded
technologies. Word has no layout engine we can query, so one-page fitting is
approximated by choosing a font preset from the content volume (the PDF is
the precisely fitted artifact)."""
import io
import re
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from .render_common import (collect_bold_terms, contact_line, cv_to_text,
                            split_segments, term_pattern)

_GREY = RGBColor(0x55, 0x55, 0x55)
_DATE_COL_CM = 3.0
_BODY_COL_CM = 21.0 - 2 * 1.25 - _DATE_COL_CM


def _preset(cv):
    """Content volume → font preset, so short CVs fill the page and long ones
    compress. (name, title, heading, body, dates)"""
    volume = len(cv_to_text(cv))
    if volume > 3200:
        return (18, 11, 10, 9.5, 8)
    if volume > 2200:
        return (19, 11.5, 10.5, 10, 8.5)
    if volume > 1400:
        return (20, 12, 11, 10.5, 9)
    return (22, 13, 12, 11.5, 9.5)


def _add_runs(paragraph, text, pattern, size, bold_all=False, color=None):
    for segment, is_bold in split_segments(text, pattern):
        run = paragraph.add_run(segment)
        run.bold = bold_all or is_bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color


def _bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '333333')
    pbdr.append(bottom)
    pPr.append(pbdr)


def _heading(doc, text, size):
    p = doc.add_paragraph()
    run = p.add_run(f'{text}:')
    run.bold = True
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    _bottom_border(p)
    return p


def _entry_table(doc, date_text, dates_size):
    """Borderless 2-column row: dates left, content right. Returns the right
    cell for the caller to fill."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(_DATE_COL_CM)
    right.width = Cm(_BODY_COL_CM)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    run = lp.add_run(date_text or '')
    run.font.size = Pt(dates_size)
    run.font.color.rgb = _GREY
    return right


def _cell_para(cell, first=True):
    if first and not cell.paragraphs[0].runs:
        return cell.paragraphs[0]
    return cell.add_paragraph()


def build_docx(cv):
    """Optimized CV dict → .docx bytes (clean reconstruction)."""
    name_sz, title_sz, heading_sz, body_sz, dates_sz = _preset(cv)
    doc = Document()
    section = doc.sections[0]
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(section, attr, Cm(1.25))
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(body_sz)
    normal.paragraph_format.space_after = Pt(2)

    pattern = term_pattern(collect_bold_terms(cv))

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv.get('name', ''))
    name_run.bold = True
    name_run.font.size = Pt(name_sz)
    name_p.paragraph_format.space_after = Pt(0)

    if cv.get('title'):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(cv['title'])
        title_run.font.size = Pt(title_sz)
        title_p.paragraph_format.space_after = Pt(1)

    contact = contact_line(cv)
    if contact:
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_p.add_run(contact)
        c_run.font.size = Pt(max(8, body_sz - 1))
        c_run.font.color.rgb = _GREY

    if cv.get('summary'):
        _heading(doc, 'Summary', heading_sz)
        p = doc.add_paragraph()
        _add_runs(p, cv['summary'], pattern, body_sz)

    if cv.get('skills_groups'):
        _heading(doc, 'Skills', heading_sz)
        for g in cv['skills_groups']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"{g.get('group', '')}: ")
            run.bold = True
            run.font.size = Pt(body_sz)
            skills_run = p.add_run(', '.join(g.get('skills') or []))
            skills_run.bold = True
            skills_run.font.size = Pt(body_sz)

    def bullets_into(cell, items):
        for bullet in items:
            p = _cell_para(cell, first=False)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.3)
            dot = p.add_run('• ')
            dot.font.size = Pt(body_sz)
            _add_runs(p, bullet, pattern, body_sz)

    if cv.get('experience'):
        _heading(doc, 'Experience', heading_sz)
        for exp in cv['experience']:
            cell = _entry_table(doc, exp.get('dates'), dates_sz)
            role = _cell_para(cell)
            role.paragraph_format.space_after = Pt(1)
            title_run = role.add_run(exp.get('title', ''))
            title_run.bold = True
            title_run.font.size = Pt(body_sz + 0.5)
            if exp.get('company'):
                comp = role.add_run(f" — {exp['company']}")
                comp.font.size = Pt(body_sz + 0.5)
            bullets_into(cell, exp.get('bullets') or [])

    if cv.get('projects'):
        _heading(doc, 'Projects', heading_sz)
        for proj in cv['projects']:
            cell = _entry_table(doc, '', dates_sz)
            head = _cell_para(cell)
            head.paragraph_format.space_after = Pt(1)
            name_run = head.add_run(proj.get('name', ''))
            name_run.bold = True
            name_run.font.size = Pt(body_sz + 0.5)
            if proj.get('tech'):
                tech = head.add_run(f" ({proj['tech']})")
                tech.font.size = Pt(max(8, body_sz - 1))
                tech.font.color.rgb = _GREY
            desc = proj.get('description') or ''
            link = proj.get('link')
            if desc or link:
                p = _cell_para(cell, first=False)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.left_indent = Cm(0.3)
                dot = p.add_run('• ')
                dot.font.size = Pt(body_sz)
                _add_runs(p, desc, pattern, body_sz)
                if link:
                    link_run = p.add_run(f'  {link}')
                    link_run.font.size = Pt(max(8, body_sz - 1))
                    link_run.font.color.rgb = _GREY

    if cv.get('education'):
        _heading(doc, 'Education', heading_sz)
        for edu in cv['education']:
            cell = _entry_table(doc, edu.get('dates'), dates_sz)
            p = _cell_para(cell)
            deg = p.add_run(edu.get('degree', ''))
            deg.bold = True
            deg.font.size = Pt(body_sz + 0.5)
            if edu.get('institution'):
                inst = p.add_run(f" — {edu['institution']}")
                inst.font.size = Pt(body_sz + 0.5)

    for extra in cv.get('extras') or []:
        if not (extra.get('lines') or []):
            continue
        _heading(doc, extra.get('heading') or 'Additional', heading_sz)
        for line in extra['lines']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.3)
            dot = p.add_run('• ')
            dot.font.size = Pt(body_sz)
            _add_runs(p, line, pattern, body_sz)

    buf = io.BytesIO()
    doc.save(buf)
    return _strip_inert_parts(buf.getvalue())


# python-docx's built-in template ships a few Word-native extras we never
# populate (customXml properties, a thumbnail image). The final package must
# contain ONLY parts we intentionally create, so they are stripped — together
# with their relationships and content-type declarations.
_STRIP_PARTS = re.compile(r'^(?:customXml/|docProps/thumbnail\.)')
_RELS_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
_CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'


def _resolve(rels_part, target):
    # OPC: <dir>/_rels/<name>.rels targets resolve relative to <dir>
    # (package root for _rels/.rels).
    base_dir = '/'.join(rels_part.split('/')[:-2])
    parts = (f'{base_dir}/{target}' if base_dir and not target.startswith('/') else target).split('/')
    out = []
    for p in parts:
        if p == '..':
            if out:
                out.pop()
        elif p not in ('', '.'):
            out.append(p)
    return '/'.join(out)


def _filter_rels(xml_bytes, rels_part):
    ET.register_namespace('', _RELS_NS)
    root = ET.fromstring(xml_bytes)
    for rel in list(root):
        target = rel.get('Target', '')
        if rel.get('TargetMode', '') != 'External' and _STRIP_PARTS.match(_resolve(rels_part, target)):
            root.remove(rel)
    return ET.tostring(root, xml_declaration=True, encoding='UTF-8')


def _filter_content_types(xml_bytes):
    ET.register_namespace('', _CT_NS)
    root = ET.fromstring(xml_bytes)
    for node in list(root):
        tag = node.tag.rsplit('}', 1)[-1]
        if tag == 'Override' and _STRIP_PARTS.match(node.get('PartName', '').lstrip('/')):
            root.remove(node)
        elif tag == 'Default' and node.get('Extension', '').lower() in ('jpeg', 'jpg', 'png'):
            root.remove(node)
    return ET.tostring(root, xml_declaration=True, encoding='UTF-8')


def _strip_inert_parts(data):
    src = zipfile.ZipFile(io.BytesIO(data))
    out = io.BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as dst:
        for name in src.namelist():
            if _STRIP_PARTS.match(name):
                continue
            payload = src.read(name)
            if name == '[Content_Types].xml':
                payload = _filter_content_types(payload)
            elif name.endswith('.rels'):
                payload = _filter_rels(payload, name)
            dst.writestr(name, payload)
    return out.getvalue()
